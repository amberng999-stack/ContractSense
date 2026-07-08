from __future__ import annotations

from io import BytesIO

from fastapi import HTTPException, UploadFile, status


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


async def read_and_validate_upload(file: UploadFile, max_upload_mb: int) -> tuple[bytes, str]:
    """
    Reads the upload once, validates size/type, and returns (content_bytes, extension).
    Callers can reuse content_bytes for multiple purposes (text extraction,
    PDF coordinate extraction, etc.) without re-reading the stream.
    """
    content = await file.read()
    max_bytes = max_upload_mb * 1024 * 1024
    if len(content) > max_bytes:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File is larger than {max_upload_mb} MB.",
        )

    extension = _extension_for(file.filename or "")
    if extension not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}.",
        )

    return content, extension


def extract_text_from_bytes(content: bytes, extension: str) -> str:
    if extension in {".txt", ".md"}:
        return _decode_text(content)
    if extension == ".pdf":
        return _extract_pdf(content)
    if extension == ".docx":
        return _extract_docx(content)
    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported file type.")


async def extract_text_from_upload(file: UploadFile, max_upload_mb: int) -> str:
    """Kept for backward compatibility — reads + extracts in one call."""
    content, extension = await read_and_validate_upload(file, max_upload_mb)
    return extract_text_from_bytes(content, extension)


def _extension_for(file_name: str) -> str:
    dot_index = file_name.rfind(".")
    if dot_index == -1:
        return ""
    return file_name[dot_index:].lower()


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "gb18030", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _extract_pdf(content: bytes) -> str:
    """
    Extracts PDF text using pdfplumber first, since it handles tables,
    multi-column layouts, and irregular spacing far more reliably than
    pypdf (pypdf frequently mangles word order/whitespace on real-world
    contract PDFs, which was the root cause of downstream LLM analysis
    reading garbled text). Falls back to pypdf only if pdfplumber fails
    or the PDF is unreadable by it (e.g. certain malformed files).
    """
    text = _extract_pdf_pdfplumber(content)
    if text.strip():
        return text
    return _extract_pdf_pypdf(content)


def _extract_pdf_pdfplumber(content: bytes) -> str:
    try:
        import pdfplumber
    except ImportError:
        return ""
    try:
        with pdfplumber.open(BytesIO(content)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(pages).strip()
    except Exception:
        # Malformed/unsupported PDF for pdfplumber — let the pypdf
        # fallback have a try instead of failing the whole request.
        return ""


def _extract_pdf_pypdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="PDF extraction dependency is not installed.") from exc

    reader = PdfReader(BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages).strip()


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="DOCX extraction dependency is not installed.") from exc

    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    ]
    return "\n".join(paragraphs + table_cells).strip()
